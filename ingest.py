"""
ingest.py
-----------------
Pipeline de ingestión multi-formato. Parsea documentos, los trocea,
genera embeddings y los almacena en ChromaDB.

Formatos soportados:
    .pdf   → PyMuPDF  (robusto con layouts complejos, columnas, tablas)
    .docx  → python-docx
    .html  → BeautifulSoup4 (extrae solo el texto, elimina tags)
    .xlsx  → openpyxl (convierte cada fila en texto narrativo)
    .txt   → lectura directa
    .md    → lectura directa

Uso:
    python ingest.py                  # procesa todo ./manuales
    python ingest.py --reset          # borra la BD y reinicia desde cero
    python ingest.py --debug          # imprime los primeros 3 chunks de cada archivo
"""

import os
os.environ["ANONYMIZED_TELEMETRY"] = "False"
# Silenciar telemetría de ChromaDB/Posthog
try:
    import posthog
    posthog.capture = lambda *args, **kwargs: None
except Exception:
    pass

import sys
import argparse
from pathlib import Path

import chromadb
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskProgressColumn
from rich.panel import Panel
from rich.table import Table
from rich.rule import Rule

from llama_index.core import VectorStoreIndex, StorageContext
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import Document
from llama_index.embeddings.ollama import OllamaEmbedding
from llama_index.vector_stores.chroma import ChromaVectorStore

import config

console = Console()

SUPPORTED_EXTS = {".pdf", ".docx", ".html", ".htm", ".xlsx", ".txt", ".md"}


# ─── Helpers ──────────────────────────────────────────────────────────────────

def get_embed_model():
    return OllamaEmbedding(
        model_name=config.EMBED_MODEL,
        base_url=config.OLLAMA_BASE_URL,
    )


def get_chroma_collection(reset: bool = False):
    client = chromadb.PersistentClient(path=config.CHROMA_PATH)
    if reset:
        try:
            client.delete_collection(config.COLLECTION_NAME)
            console.print("[yellow]⚠  Colección anterior eliminada.[/yellow]")
        except Exception:
            pass
    collection = client.get_or_create_collection(config.COLLECTION_NAME)
    return client, collection


# ─── Parsers por formato ───────────────────────────────────────────────────────

def parse_pdf(path: Path) -> list:
    """PyMuPDF: una página = un Document. Preserva layout, filtra líneas vacías."""
    try:
        import fitz
    except ImportError:
        console.print("[red]✗ PyMuPDF no instalado. Ejecuta: pip install pymupdf[/red]")
        sys.exit(1)

    docs = []
    try:
        doc = fitz.open(str(path))
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            lines = [l for l in text.split("\n") if len(l.strip()) > 3]
            text = "\n".join(lines).strip()
            if not text:
                continue
            docs.append(Document(
                text=text,
                metadata={
                    "file_name":   path.name,
                    "file_path":   str(path),
                    "page_label":  str(page_num + 1),
                    "page":        page_num + 1,
                    "total_pages": len(doc),
                    "source":      "pymupdf",
                    "format":      "pdf",
                }
            ))
        doc.close()
    except Exception as e:
        console.print(f"  [red]✗[/red] {path.name} — Error PDF: {e}")
    return docs


def parse_docx(path: Path) -> list:
    """python-docx: extrae párrafos no vacíos como un único Document."""
    try:
        import docx
    except ImportError:
        console.print("[red]✗ python-docx no instalado. Ejecuta: pip install python-docx[/red]")
        return []
    try:
        doc = docx.Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)
        if not text:
            return []
        return [Document(
            text=text,
            metadata={"file_name": path.name, "file_path": str(path), "format": "docx"}
        )]
    except Exception as e:
        console.print(f"  [red]✗[/red] {path.name} — Error DOCX: {e}")
        return []


def parse_html(path: Path) -> list:
    """BeautifulSoup: elimina tags y scripts, extrae texto limpio."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        console.print("[red]✗ beautifulsoup4 no instalado. Ejecuta: pip install beautifulsoup4[/red]")
        return []
    try:
        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "header", "footer", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        lines = [l.strip() for l in text.splitlines() if len(l.strip()) > 3]
        text = "\n".join(lines)
        if not text:
            return []
        return [Document(
            text=text,
            metadata={"file_name": path.name, "file_path": str(path), "format": "html"}
        )]
    except Exception as e:
        console.print(f"  [red]✗[/red] {path.name} — Error HTML: {e}")
        return []


def parse_xlsx(path: Path) -> list:
    """
    openpyxl: convierte cada hoja en texto narrativo.
    Cada fila se representa como 'Columna: valor' para preservar contexto semántico.
    """
    try:
        import openpyxl
    except ImportError:
        console.print("[red]✗ openpyxl no instalado. Ejecuta: pip install openpyxl[/red]")
        return []
    docs = []
    try:
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            headers = [str(h).strip() if h is not None else f"Col{i}"
                       for i, h in enumerate(rows[0])]
            text_blocks = []
            for row in rows[1:]:
                parts = [f"{header}: {str(cell).strip()}"
                         for header, cell in zip(headers, row)
                         if cell is not None and str(cell).strip()]
                if parts:
                    text_blocks.append(" | ".join(parts))
            if not text_blocks:
                continue
            docs.append(Document(
                text="\n".join(text_blocks),
                metadata={
                    "file_name":  path.name,
                    "file_path":  str(path),
                    "sheet_name": sheet_name,
                    "format":     "xlsx",
                }
            ))
        wb.close()
    except Exception as e:
        console.print(f"  [red]✗[/red] {path.name} — Error XLSX: {e}")
    return docs


def parse_text(path: Path) -> list:
    """TXT y MD: lectura directa."""
    try:
        text = path.read_text(encoding="utf-8", errors="ignore").strip()
        if not text:
            return []
        return [Document(
            text=text,
            metadata={
                "file_name": path.name,
                "file_path": str(path),
                "format":    path.suffix.lstrip("."),
            }
        )]
    except Exception as e:
        console.print(f"  [red]✗[/red] {path.name} — Error TXT/MD: {e}")
        return []


# ─── Router ───────────────────────────────────────────────────────────────────

PARSERS = {
    ".pdf":  parse_pdf,
    ".docx": parse_docx,
    ".html": parse_html,
    ".htm":  parse_html,
    ".xlsx": parse_xlsx,
    ".txt":  parse_text,
    ".md":   parse_text,
}


def load_documents() -> list:
    docs_path = Path(config.DOCS_DIR)
    if not docs_path.exists() or not any(docs_path.iterdir()):
        console.print(f"[red]✗ No hay documentos en '{config.DOCS_DIR}'.[/red]")
        console.print("  Coloca archivos PDF, DOCX, HTML, XLSX, TXT o MD en esa carpeta.")
        sys.exit(1)

    all_files = [
        f for f in sorted(docs_path.rglob("*"))
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTS
    ]

    if not all_files:
        console.print(f"[red]✗ No se encontraron archivos soportados en '{config.DOCS_DIR}'.[/red]")
        console.print(f"  Formatos: {', '.join(sorted(SUPPORTED_EXTS))}")
        sys.exit(1)

    by_ext: dict = {}
    for f in all_files:
        by_ext.setdefault(f.suffix.lower(), []).append(f)

    summary = "  ".join(
        f"[cyan]{ext}[/cyan]×{len(files)}" for ext, files in sorted(by_ext.items())
    )
    console.print(f"[cyan]📂 {len(all_files)} archivo(s) encontrado(s):[/cyan]  {summary}\n")

    documents = []
    for file_path in all_files:
        ext = file_path.suffix.lower()
        docs = PARSERS[ext](file_path)
        if docs:
            documents.extend(docs)
            label = f"{len(docs)} página(s)" if ext == ".pdf" else "ok"
            console.print(f"  [green]✓[/green] {file_path.name}  [dim]({label})[/dim]")

    console.print()
    return documents


# ─── Split ────────────────────────────────────────────────────────────────────

def split_documents(documents: list):
    splitter = SentenceSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
    )
    return splitter.get_nodes_from_documents(documents, show_progress=False)


# ─── Debug ────────────────────────────────────────────────────────────────────

def print_debug(nodes, n: int = 3):
    console.print(Rule("[yellow]DEBUG — primeros chunks[/yellow]"))
    for i, node in enumerate(nodes[:n]):
        meta = node.metadata or {}
        console.print(
            f"\n[bold yellow]Chunk {i}[/bold yellow]  "
            f"[dim]{meta.get('file_name','?')} · "
            f"pág. {meta.get('page_label', meta.get('sheet_name', '—'))}[/dim]"
        )
        console.print(node.text[:600])
        console.print()
    console.print(Rule())


# ─── Resumen ──────────────────────────────────────────────────────────────────

def print_summary(documents, nodes):
    table = Table(title="Resumen de ingestión", show_header=True, header_style="bold cyan")
    table.add_column("Métrica", style="dim")
    table.add_column("Valor", justify="right")

    files   = {doc.metadata.get("file_name", "?") for doc in documents}
    formats = {doc.metadata.get("format",    "?") for doc in documents}

    table.add_row("Archivos únicos",          str(len(files)))
    table.add_row("Formatos",                 ", ".join(sorted(formats)))
    table.add_row("Páginas / bloques leídos", str(len(documents)))
    table.add_row("Chunks generados",         str(len(nodes)))
    table.add_row("Chunk size",               str(config.CHUNK_SIZE))
    table.add_row("Chunk overlap",            str(config.CHUNK_OVERLAP))
    table.add_row("Parser PDF",               "PyMuPDF")
    table.add_row("Modelo embeddings",        config.EMBED_MODEL)
    table.add_row("LLM",                      config.LLM_MODEL)

    console.print(table)
    console.print()
    console.print("[bold]Archivos procesados:[/bold]")
    for f in sorted(files):
        console.print(f"  [green]✓[/green] {f}")


# ─── Pipeline principal ───────────────────────────────────────────────────────

def ingest(reset: bool = False, debug: bool = False):
    console.print(Panel.fit(
        "[bold cyan]RAG Manuales — Ingestión (PyMuPDF)[/bold cyan]",
        border_style="cyan"
    ))

    # 1. Verificar Ollama
    console.print("[cyan]🔌 Verificando conexión con Ollama...[/cyan]")
    try:
        embed_model = get_embed_model()
        embed_model.get_text_embedding("test")
        console.print(f"  [green]✓[/green] Ollama OK  ({config.EMBED_MODEL})\n")
    except Exception as e:
        console.print(f"[red]✗ No se puede conectar con Ollama: {e}[/red]")
        console.print("  Asegúrate de que Ollama está corriendo y de haber ejecutado:")
        console.print(f"  [bold]ollama pull {config.EMBED_MODEL}[/bold]")
        sys.exit(1)

    # 2. Cargar
    documents = load_documents()

    # 3. Trocear
    console.print(f"[cyan]✂  Troceando (size={config.CHUNK_SIZE}, overlap={config.CHUNK_OVERLAP})...[/cyan]")
    nodes = split_documents(documents)
    console.print(f"  [green]✓[/green] {len(nodes)} chunks listos\n")

    # 4. Debug opcional
    if debug:
        print_debug(nodes)

    # 5. ChromaDB
    _, collection   = get_chroma_collection(reset=reset)
    vector_store    = ChromaVectorStore(chroma_collection=collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)

    # 6. Embeddings + guardado en lotes
    console.print("[cyan]🧮 Generando embeddings y almacenando en ChromaDB...[/cyan]")
    console.print("  (puede tardar varios minutos según el volumen)\n")

    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TaskProgressColumn(),
        console=console,
    ) as progress:
        task = progress.add_task("Vectorizando chunks...", total=len(nodes))
        batch_size = 20
        for i in range(0, len(nodes), batch_size):
            batch = nodes[i : i + batch_size]
            VectorStoreIndex(
                batch,
                storage_context=storage_context,
                embed_model=embed_model,
            )
            progress.update(task, advance=len(batch))

    console.print()
    console.print("[bold green]✅ Ingestión completada.[/bold green]")
    console.print(f"   Base de datos guardada en: [bold]{config.CHROMA_PATH}[/bold]\n")

    print_summary(documents, nodes)


# ─── Entry point ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Ingestar manuales en ChromaDB (multi-formato)")
    parser.add_argument("--reset", action="store_true",
                        help="Borrar la BD existente antes de ingestar")
    parser.add_argument("--debug", action="store_true",
                        help="Imprimir los primeros 3 chunks para inspeccionar el parseo")
    args = parser.parse_args()
    ingest(reset=args.reset, debug=args.debug)
